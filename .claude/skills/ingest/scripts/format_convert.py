#!/usr/bin/env python3
"""
格式转换工具

将 Word (.docx) 或其他格式转换为标准的 Markdown 文本，
供 AI 提取流程使用。

用法:
    python format_convert.py <input.{docx,html}> <output.md>

依赖:
    python-docx (pip install python-docx)

完整规格见 DEV_SPEC 8.1.1 输入类型与处理策略
"""

import sys
from pathlib import Path


def docx_to_markdown(input_path: str) -> str:
    """
    将 .docx 转换为 Markdown 文本

    Args:
        input_path: docx 文件路径

    Returns:
        Markdown 格式文本
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed. Run: pip install python-docx"
        )

    doc = Document(input_path)
    lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        # 简单根据样式判断标题层级
        style_name = para.style.name.lower() if para.style else ""
        if "heading" in style_name or "标题" in style_name:
            level = 1
            for i in range(1, 10):
                if f"heading {i}" in style_name or f"标题 {i}" in style_name:
                    level = i
                    break
            lines.append(f"{'#' * level} {text}")
        else:
            lines.append(text)

    # 简单处理表格
    for table in doc.tables:
        lines.append("")
        # 表头
        if table.rows:
            header_cells = [cell.text.strip() for cell in table.rows[0].cells]
            lines.append("| " + " | ".join(header_cells) + " |")
            lines.append("| " + " | ".join(["---"] * len(header_cells)) + " |")
            for row in table.rows[1:]:
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


def convert(input_path: str, output_path: str | None = None) -> str:
    """
    通用格式转换入口

    Args:
        input_path: 输入文件路径
        output_path: 输出文件路径，None 则返回文本

    Returns:
        若 output_path 为 None 返回文本，否则返回输出路径
    """
    inp = Path(input_path)
    if not inp.exists():
        raise FileNotFoundError(f"Input not found: {input_path}")

    suffix = inp.suffix.lower()

    if suffix == ".docx":
        md = docx_to_markdown(input_path)
    elif suffix in (".md", ".markdown", ".txt"):
        with open(input_path, "r", encoding="utf-8") as f:
            md = f.read()
    elif suffix in (".html", ".htm"):
        try:
            from bs4 import BeautifulSoup
            with open(input_path, "r", encoding="utf-8") as f:
                soup = BeautifulSoup(f.read(), "html.parser")
            # 移除 script/style
            for tag in soup(["script", "style"]):
                tag.decompose()
            md = soup.get_text(separator="\n", strip=True)
        except ImportError:
            raise ImportError("beautifulsoup4 not installed. Run: pip install beautifulsoup4")
    else:
        raise ValueError(f"Unsupported format: {suffix}")

    if output_path:
        out = Path(output_path)
        out.parent.mkdir(parents=True, exist_ok=True)
        with open(out, "w", encoding="utf-8") as f:
            f.write(md)
        return str(out)
    return md


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python format_convert.py <input> [output.md]", file=sys.stderr)
        sys.exit(1)

    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    try:
        result = convert(inp, out)
        if out:
            print(f"[OK] Converted to {result}")
        else:
            print(result)
    except Exception as e:
        print(f"[ERROR] {e}", file=sys.stderr)
        sys.exit(1)
